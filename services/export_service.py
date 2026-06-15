import os
import io
import json
import csv
import markdown
from models import get_db
from datetime import datetime

EXPORT_DIR = 'exports'
os.makedirs(EXPORT_DIR, exist_ok=True)

def export_review_markdown(review_id):
    db = get_db()
    review = db.execute('SELECT * FROM literature_reviews WHERE id=?', (review_id,)).fetchone()
    db.close()
    if not review:
        return None
    review = dict(review)
    md = f"# {review['title']}\n\n"
    md += f"**Generated:** {review.get('created_at', '')}\n\n"
    md += "---\n\n"
    if review.get('introduction'):
        md += f"## Introduction\n\n{review['introduction']}\n\n"
    if review.get('thematic_analysis'):
        md += f"## Thematic Analysis\n\n{review['thematic_analysis']}\n\n"
    if review.get('research_gaps'):
        md += f"## Research Gaps\n\n{review['research_gaps']}\n\n"
    if review.get('future_directions'):
        md += f"## Future Directions\n\n{review['future_directions']}\n\n"
    if review.get('conclusion'):
        md += f"## Conclusion\n\n{review['conclusion']}\n\n"
    filename = f"review_{review_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
    filepath = os.path.join(EXPORT_DIR, filename)
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(md)
    _save_export_record(review_id, filename, filepath, 'markdown')
    return filepath

def export_review_html(review_id):
    db = get_db()
    review = db.execute('SELECT * FROM literature_reviews WHERE id=?', (review_id,)).fetchone()
    db.close()
    if not review:
        return None
    review = dict(review)
    md_content = ''
    if review.get('introduction'):
        md_content += f"## Introduction\n\n{review['introduction']}\n\n"
    if review.get('thematic_analysis'):
        md_content += f"## Thematic Analysis\n\n{review['thematic_analysis']}\n\n"
    if review.get('research_gaps'):
        md_content += f"## Research Gaps\n\n{review['research_gaps']}\n\n"
    if review.get('future_directions'):
        md_content += f"## Future Directions\n\n{review['future_directions']}\n\n"
    if review.get('conclusion'):
        md_content += f"## Conclusion\n\n{review['conclusion']}\n\n"
    html_body = markdown.markdown(md_content)
    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>{review['title']}</title>
<style>
body {{ font-family: Georgia, serif; max-width: 800px; margin: 40px auto; line-height: 1.8; color: #333; }}
h1, h2 {{ color: #1a1a2e; }} h2 {{ border-bottom: 2px solid #2563eb; padding-bottom: 8px; }}
p {{ text-align: justify; }} .meta {{ color: #666; font-style: italic; }}
</style>
</head>
<body>
<h1>{review['title']}</h1>
<p class="meta">Generated: {review.get('created_at','')}</p>
{html_body}
</body>
</html>"""
    filename = f"review_{review_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
    filepath = os.path.join(EXPORT_DIR, filename)
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(html)
    _save_export_record(review_id, filename, filepath, 'html')
    return filepath

def export_review_docx(review_id):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return None
    db = get_db()
    review = db.execute('SELECT * FROM literature_reviews WHERE id=?', (review_id,)).fetchone()
    db.close()
    if not review:
        return None
    review = dict(review)
    doc = Document()
    title = doc.add_heading(review['title'], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated: {review.get('created_at', '')}")
    sections_map = [
        ('introduction', 'Introduction'),
        ('thematic_analysis', 'Thematic Analysis'),
        ('research_gaps', 'Research Gaps'),
        ('future_directions', 'Future Directions'),
        ('conclusion', 'Conclusion'),
    ]
    for key, heading in sections_map:
        if review.get(key):
            doc.add_heading(heading, level=1)
            doc.add_paragraph(review[key])
    filename = f"review_{review_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(EXPORT_DIR, filename)
    doc.save(filepath)
    _save_export_record(review_id, filename, filepath, 'docx')
    return filepath

def export_literature_matrix_csv(paper_ids=None):
    db = get_db()
    if paper_ids:
        placeholders = ','.join('?' * len(paper_ids))
        papers = db.execute(f'SELECT * FROM papers WHERE id IN ({placeholders})', paper_ids).fetchall()
    else:
        papers = db.execute('SELECT * FROM papers ORDER BY year DESC').fetchall()
    db.close()
    filename = f"matrix_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
    filepath = os.path.join(EXPORT_DIR, filename)
    with open(filepath, 'w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        writer.writerow(['Title', 'Authors', 'Year', 'Journal', 'DOI', 'Abstract', 'Citations', 'Source', 'Keywords'])
        for p in papers:
            writer.writerow([p['title'], p['authors'], p['year'], p['journal'],
                             p['doi'], (p['abstract'] or '')[:200], p['citations_count'], p['source'], p['keywords']])
    _save_export_record(None, filename, filepath, 'csv')
    return filepath

def export_literature_matrix_excel(paper_ids=None):
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment
    except ImportError:
        return export_literature_matrix_csv(paper_ids)
    db = get_db()
    if paper_ids:
        placeholders = ','.join('?' * len(paper_ids))
        papers = db.execute(f'SELECT * FROM papers WHERE id IN ({placeholders})', paper_ids).fetchall()
    else:
        papers = db.execute('SELECT * FROM papers ORDER BY year DESC').fetchall()
    db.close()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Literature Matrix'
    headers = ['Title', 'Authors', 'Year', 'Journal', 'DOI', 'Abstract', 'Citations', 'Source', 'Keywords']
    header_fill = PatternFill(start_color='2563EB', end_color='2563EB', fill_type='solid')
    header_font = Font(bold=True, color='FFFFFF')
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center')
    for row_idx, p in enumerate(papers, 2):
        ws.cell(row=row_idx, column=1, value=p['title'])
        ws.cell(row=row_idx, column=2, value=p['authors'])
        ws.cell(row=row_idx, column=3, value=p['year'])
        ws.cell(row=row_idx, column=4, value=p['journal'])
        ws.cell(row=row_idx, column=5, value=p['doi'])
        ws.cell(row=row_idx, column=6, value=(p['abstract'] or '')[:200])
        ws.cell(row=row_idx, column=7, value=p['citations_count'])
        ws.cell(row=row_idx, column=8, value=p['source'])
        ws.cell(row=row_idx, column=9, value=p['keywords'])
    filename = f"matrix_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    filepath = os.path.join(EXPORT_DIR, filename)
    wb.save(filepath)
    _save_export_record(None, filename, filepath, 'excel')
    return filepath

def export_bibtex_all(paper_ids=None):
    from services.citation_service import generate_bibtex
    db = get_db()
    if paper_ids:
        placeholders = ','.join('?' * len(paper_ids))
        papers = db.execute(f'SELECT * FROM papers WHERE id IN ({placeholders})', paper_ids).fetchall()
    else:
        papers = db.execute('SELECT * FROM papers').fetchall()
    db.close()
    bibtex_content = '\n\n'.join([generate_bibtex(dict(p)) for p in papers])
    filename = f"references_{datetime.now().strftime('%Y%m%d_%H%M%S')}.bib"
    filepath = os.path.join(EXPORT_DIR, filename)
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(bibtex_content)
    _save_export_record(None, filename, filepath, 'bibtex')
    return filepath

def export_ris_all(paper_ids=None):
    from services.citation_service import generate_ris
    db = get_db()
    if paper_ids:
        placeholders = ','.join('?' * len(paper_ids))
        papers = db.execute(f'SELECT * FROM papers WHERE id IN ({placeholders})', paper_ids).fetchall()
    else:
        papers = db.execute('SELECT * FROM papers').fetchall()
    db.close()
    ris_content = '\n'.join([generate_ris(dict(p)) for p in papers])
    filename = f"references_{datetime.now().strftime('%Y%m%d_%H%M%S')}.ris"
    filepath = os.path.join(EXPORT_DIR, filename)
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(ris_content)
    _save_export_record(None, filename, filepath, 'ris')
    return filepath

def _save_export_record(ref_id, filename, filepath, fmt):
    try:
        db = get_db()
        db.execute('INSERT INTO exports (type, filename, file_path, format, status) VALUES (?, ?, ?, ?, ?)',
                   ('review' if ref_id else 'general', filename, filepath, fmt, 'ready'))
        db.execute('INSERT INTO notifications (type, title, message) VALUES (?, ?, ?)',
                   ('export', 'Export Ready', f'{filename} is ready for download.'))
        db.commit()
        db.close()
    except Exception as e:
        try:
            db2 = get_db()
            db2.execute('INSERT INTO system_logs (level, source, message) VALUES (?, ?, ?)',
                        ('WARN', 'export_service:save_record', str(e)[:500]))
            db2.commit()
            db2.close()
        except Exception:
            pass
